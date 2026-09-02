"""Word 财务报告解析器。

运行环境：Python 3.11，Conda 环境 ``rag_311``。
依赖：``python-docx>=1.1``、``pandas>=2.0``。

数据流向：
    .docx -> Document -> 按文档顺序遍历 Paragraph/Table
          -> 标题层级跟踪 + 表格清洗
          -> List[Dict] + List[pandas.DataFrame]
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterator, TypeAlias

import pandas as pd
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

PathLike: TypeAlias = str | Path
ParagraphRecord: TypeAlias = dict[str, Any]

LOGGER = logging.getLogger(__name__)
_WHITESPACE_RE = re.compile(r"\s+")
_HEADING_STYLE_RE = re.compile(r"(?:heading|标题)\s*([1-9]\d*)", re.IGNORECASE)


class WordParseError(RuntimeError):
    """Word 文档结构无法正常解析。"""


@dataclass(slots=True)
class WordParseResult:
    """Word 解析结果，段落与表格均带有标题上下文。"""

    paragraphs: list[ParagraphRecord]
    tables: list[pd.DataFrame]

    @property
    def text(self) -> str:
        """返回可直接用于切块或全文检索的纯文本。"""

        return "\n".join(record["text"] for record in self.paragraphs)

    @property
    def paragraph_texts(self) -> list[str]:
        """仅返回段落文本，用于兼容不需要层级信息的业务。"""

        return [record["text"] for record in self.paragraphs]

    @property
    def table_records(self) -> list[list[dict[str, Any]]]:
        """将所有表格转成适合 JSON 序列化的 List[Dict]。"""

        return [
            [
                {str(key): value for key, value in record.items()}
                for record in (
                    frame.astype(object)
                    .where(frame.notna(), None)
                    .to_dict(orient="records")
                )
            ]
            for frame in self.tables
        ]



def _validate_docx_path(file_path: PathLike) -> Path:
    """验证输入路径，在调用第三方库前给出明确错误。"""

    path = Path(file_path).expanduser().resolve()
    if path.suffix.lower() != ".docx":
        raise ValueError(f"仅支持 .docx 文件：{path}")
    if not path.is_file():
        raise FileNotFoundError(f"Word 文件不存在：{path}")
    return path


def _clean_text(value: str | None) -> str:
    """合并换行、制表符和重复空格，保留实际文本内容。"""

    return _WHITESPACE_RE.sub(" ", value or "").strip()


def _iter_blocks(document: DocxDocument) -> Iterator[Paragraph | Table]:
    """按 Word XML 中的真实顺序产出段落和表格。

    ``document.paragraphs`` 和 ``document.tables`` 是两个独立集合，分别遍历
    会丢失“段落 -> 表格 -> 段落”的原始次序。因此直接遍历
    ``document.element.body``，再将 XML 节点包装为 python-docx 对象。
    """

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _heading_level(paragraph: Paragraph) -> int | None:
    """从段落样式或 XML outlineLvl 推断标题级别。"""

    style_name = paragraph.style.name if paragraph.style is not None else ""
    style_name = style_name or ""
    match = _HEADING_STYLE_RE.search(style_name)

    if match:
        return int(match.group(1))
    if style_name.casefold() in {"title", "标题"}:
        return 1

    paragraph_properties = paragraph._p.pPr  # noqa: SLF001 - python-docx 未暴露大纲级别 API
    if paragraph_properties is not None:
        outline = paragraph_properties.find(qn("w:outlineLvl"))
        if outline is not None:
            value = outline.get(qn("w:val"))
            if value is not None and value.isdigit():
                return int(value) + 1
    return None


def _update_heading_path(path: list[str], level: int, title: str) -> list[str]:
    """将当前标题写入层级路径，并移除已经结束的子级标题。"""

    level = max(level, 1)
    parent_path = path[: level - 1]
    if len(parent_path) < level - 1:
        # 某些文档从三级标题开始，用空值保持级别位置。
        parent_path.extend([""] * (level - 1 - len(parent_path)))
    return [*parent_path, title]


def _remove_empty_columns(rows: list[list[str]]) -> list[list[str]]:
    """统一行宽并删除整列为空的 Word 占位列。"""

    if not rows:
        return []
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    keep = [index for index in range(width) if any(row[index] for row in normalized)]
    return [[row[index] for index in keep] for row in normalized]


def _build_headers(header_rows: list[list[str]], width: int) -> list[str]:
    """合并多行表头，并生成唯一且非空的列名。"""

    raw_headers: list[str] = []
    for column_index in range(width):
        parts: list[str] = []
        for row in header_rows:
            value = row[column_index]
            if value and (not parts or value != parts[-1]):
                parts.append(value)
        raw_headers.append(" / ".join(parts) or f"column_{column_index + 1}")

    counts: dict[str, int] = {}
    unique_headers: list[str] = []
    for header in raw_headers:
        counts[header] = counts.get(header, 0) + 1
        suffix = f"_{counts[header]}" if counts[header] > 1 else ""
        unique_headers.append(f"{header}{suffix}")
    return unique_headers


def table_to_dataframe(
    table: Table,
    *,
    header_rows: int = 1,  # 🔧【可调参数】复合表头占用的行数，0 表示无表头
) -> pd.DataFrame:
    """清洗 Word 表格并转为 DataFrame。

    python-docx 的 ``Table -> Row -> Cell`` 是对 WordprocessingML 表格网格的
    高层封装。这里先将 Cell 文本规范化，再统一列数、删除空列，
    最后才构建 DataFrame，避免不规则表格导致列偏移。
    """

    if header_rows < 0:
        raise ValueError("header_rows 不能小于 0")

    rows = [[_clean_text(cell.text) for cell in row.cells] for row in table.rows]
    rows = _remove_empty_columns([row for row in rows if any(row)])
    if not rows:
        return pd.DataFrame()
    if header_rows >= len(rows) and header_rows != 0:
        LOGGER.warning("表头行数 %s 不小于表格总行数 %s", header_rows, len(rows))

    width = len(rows[0])
    if header_rows == 0:
        columns = [f"column_{index + 1}" for index in range(width)]
        data_rows = rows
    else:
        actual_header_rows = min(header_rows, len(rows))
        columns = _build_headers(rows[:actual_header_rows], width)
        data_rows = rows[actual_header_rows:]

    frame = pd.DataFrame(data_rows, columns=columns)
    frame.replace({"": pd.NA}, inplace=True)

    # 💡【面试加分点】财务表格中的行列关系表达“指标-期间-金额”语义。
    # 若直接拼成纯文本，检索后很难准确计算同比、比率或跨期对比；
    # DataFrame 则保留了可计算的结构，后续 Agent 可使用确定性工具而非让 LLM 心算。
    return frame


def parse_word(
    file_path: PathLike,
    *,
    header_rows: int = 1,  # 🔧【可调参数】Word 表格的默认表头行数
    include_empty_paragraphs: bool = False,  # 🔧【可调参数】是否保留空段落
) -> WordParseResult:
    """解析 .docx，返回带标题层级的段落字典和结构化表格。"""

    path = _validate_docx_path(file_path)
    try:
        document = Document(str(path))
    except Exception as exc:
        raise WordParseError(f"Word 文档打开失败：{path}") from exc

    paragraphs: list[ParagraphRecord] = []
    tables: list[pd.DataFrame] = []
    heading_path: list[str] = []
    paragraph_index = 0

    try:
        for block_index, block in enumerate(_iter_blocks(document)):
            if isinstance(block, Paragraph):
                text = _clean_text(block.text)
                if not text and not include_empty_paragraphs:
                    continue

                level = _heading_level(block) if text else None
                if level is not None:
                    heading_path = _update_heading_path(heading_path, level, text)

                paragraphs.append(
                    {
                        "index": paragraph_index,
                        "block_index": block_index,
                        "text": text,
                        "style": block.style.name if block.style is not None else None,
                        "is_heading": level is not None,
                        "heading_level": level,
                        "heading_path": [title for title in heading_path if title],
                    }
                )
                paragraph_index += 1
                continue

            frame = table_to_dataframe(block, header_rows=header_rows)
            frame.attrs.update(
                {
                    "table_index": len(tables),
                    "block_index": block_index,
                    "heading_path": [title for title in heading_path if title],
                }
            )
            tables.append(frame)
    except Exception as exc:
        if isinstance(exc, (ValueError, WordParseError)):
            raise
        raise WordParseError(f"Word 文档内容解析失败：{path}") from exc

    LOGGER.info(
        "Word 解析完成：file=%s, paragraphs=%s, tables=%s",
        path,
        len(paragraphs),
        len(tables),
    )
    return WordParseResult(paragraphs=paragraphs, tables=tables)


__all__ = [
    "WordParseError",
    "WordParseResult",
    "parse_word",
    "table_to_dataframe",
]
